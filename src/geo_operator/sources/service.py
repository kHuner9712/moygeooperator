from __future__ import annotations

import io
import json
import mimetypes
import re
import uuid
from pathlib import Path
from typing import Any, ClassVar

from docx import Document
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader

from geo_operator.core.db import Database
from geo_operator.core.storage import ArtifactStore
from geo_operator.core.time import utc_now


class SourceIngestionService:
    MAX_BYTES = 50 * 1024 * 1024
    ALLOWED_EXTENSIONS: ClassVar[set[str]] = {
        ".txt",
        ".md",
        ".csv",
        ".json",
        ".pdf",
        ".docx",
        ".xlsx",
        ".xlsm",
        ".png",
        ".jpg",
        ".jpeg",
        ".webp",
        ".gif",
    }

    def __init__(self, database: Database, artifacts: ArtifactStore) -> None:
        self.database, self.artifacts = database, artifacts

    def ingest(
        self,
        tenant_id: str,
        original_name: str,
        content: bytes,
        media_type: str | None = None,
    ) -> dict[str, Any]:
        if not self.database.one(
            "SELECT id FROM tenants WHERE id=? AND status='ACTIVE'", (tenant_id,)
        ):
            raise KeyError("Tenant not found")
        if not content:
            raise ValueError("Source file cannot be empty")
        if len(content) > self.MAX_BYTES:
            raise ValueError("Source file exceeds the 50 MiB limit")
        safe_name = self._safe_name(original_name)
        extension = Path(safe_name).suffix.lower()
        if extension not in self.ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported source file type: {extension or 'none'}")
        asset_id = uuid.uuid4().hex
        relative = f"source/{asset_id}/{safe_name}"
        extracted, metadata, status = self._extract(extension, content)
        _, content_hash = self.artifacts.atomic_write(tenant_id, relative, content)
        text_relative = None
        if extracted is not None:
            text_relative = f"source/extracted/{asset_id}.txt"
            self.artifacts.atomic_write(tenant_id, text_relative, extracted.encode("utf-8"))
        now = utc_now()
        resolved_media_type = (
            media_type or mimetypes.guess_type(safe_name)[0] or "application/octet-stream"
        )
        with self.database.transaction() as connection:
            connection.execute(
                """INSERT INTO source_assets(
                   id,tenant_id,original_name,media_type,relative_path,extracted_text_path,
                   content_sha256,size,extraction_status,metadata_json,created_at)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
                (
                    asset_id,
                    tenant_id,
                    safe_name,
                    resolved_media_type,
                    relative,
                    text_relative,
                    content_hash,
                    len(content),
                    status,
                    json.dumps(metadata, ensure_ascii=False),
                    now,
                ),
            )
        return self.get(asset_id)

    def list(self, tenant_id: str) -> list[dict[str, Any]]:
        return self.database.all(
            "SELECT * FROM source_assets WHERE tenant_id=? ORDER BY created_at", (tenant_id,)
        )

    def get(self, asset_id: str) -> dict[str, Any]:
        row = self.database.one("SELECT * FROM source_assets WHERE id=?", (asset_id,))
        if not row:
            raise KeyError("Source asset not found")
        row["metadata"] = json.loads(str(row["metadata_json"]))
        return row

    @staticmethod
    def _safe_name(name: str) -> str:
        name = Path(name or "").name.strip()
        name = re.sub(r"[^0-9A-Za-z._()\-\u4e00-\u9fff ]+", "_", name)
        if not name or name in {".", ".."}:
            raise ValueError("A valid source filename is required")
        return name[:240]

    @staticmethod
    def _extract(extension: str, content: bytes) -> tuple[str | None, dict[str, Any], str]:
        if extension in {".txt", ".md", ".csv", ".json"}:
            return content.decode("utf-8-sig"), {}, "EXTRACTED"
        if extension == ".pdf":
            reader = PdfReader(io.BytesIO(content))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            return text, {"pages": len(reader.pages)}, "EXTRACTED"
        if extension == ".docx":
            document = Document(io.BytesIO(content))
            blocks = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                blocks.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
            return "\n".join(blocks), {"paragraphs": len(document.paragraphs)}, "EXTRACTED"
        if extension in {".xlsx", ".xlsm"}:
            workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            lines: list[str] = []
            for sheet in workbook.worksheets:
                lines.append(f"# {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    lines.append("\t".join("" if value is None else str(value) for value in row))
            return "\n".join(lines), {"sheets": workbook.sheetnames}, "EXTRACTED"
        with Image.open(io.BytesIO(content)) as image:
            metadata = {"width": image.width, "height": image.height, "format": image.format}
        return None, metadata, "STORED_NO_TEXT"
